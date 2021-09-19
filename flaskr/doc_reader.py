
import docx
from . import doc_analyzer
from . import pptxcreator

def final_create(theme,filename):
    doc = docx.Document(filename)
    print(len(doc.paragraphs))
    paragraphs = []
    title_headings = []
    headings = []
    for i,par in enumerate(doc.paragraphs):
        
        if len(par.text) >  60:    
            paragraphs.append(par)
    for i,par in enumerate(paragraphs):
        
        
        a = doc_analyzer.summarize_heading(par.text)
        
        headings.append(a)
        title_headings.append("Page " + str(i))
        
    filename_new = filename[:filename.find('.')].title()
    
    pptxcreator.create_slides(theme,filename_new,title_headings,headings)

